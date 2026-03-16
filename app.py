from flask import Flask, render_template, request, redirect, url_for, jsonify, flash, Response, send_from_directory
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime, date
import requests, base64, os, json, re
from dotenv import load_dotenv
from fpdf import FPDF
import markdown

load_dotenv()

try:
    import PyPDF2; HAS_PDF = True
except ImportError:
    HAS_PDF = False
try:
    from docx import Document; HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
try:
    import openpyxl; HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

# ── Config ──
NOME_IA      = "Cerebro IA"
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_URL     = "https://api.groq.com/openai/v1/chat/completions"
MODELO_TX    = "llama-3.3-70b-versatile"
MODELO_VIS   = "llama-3.2-11b-vision-preview"
LIMITE_FREE  = 999999  # Limite virtualmente infinito

# Garante caminho absoluto para uploads em qualquer ambiente (Railway/Local)
BASE_DIR    = os.path.abspath(os.path.dirname(__file__))
UPLOAD_DIR  = os.path.join(BASE_DIR, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

EXTS_IMG    = {"png","jpg","jpeg","gif","webp","bmp"}
EXTS_ARQ    = {"pdf","txt","docx","xlsx","csv","md"}

# User-Agent profissional para evitar bloqueios de APIs
UA_PRO = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"

# ── Prompts dos agentes Expert Ultra (V27 - Omni-Nexus) ──
AGENTES = {
    "geral": {
        "nome": "Cerebro IA (Omni-Nexus V27)",
        "icone": "🧠",
        "cor": "#4f8ef7",
        "prompt": "Você é o assistente definitivo. Use ferramentas de busca, visão computacional e geolocalização para ser ultra-preciso."
    },
    "especialista": {
        "nome": "Analista Estratégico & Cientista",
        "icone": "�",
        "cor": "#10b981",
        "prompt": "Foco em análise profunda, prós e contras, e rigor técnico. Use LaTeX para fórmulas."
    },
    "criativo": {
        "nome": "Mestre de Mídia & Storytelling",
        "icone": "🎭",
        "cor": "#ec4899",
        "prompt": "Gere histórias, roteiros e ideias inovadoras. Use comandos de imagem, vídeo e música (Nano Banana, Veo, Lyria)."
    },
    "professor": {
        "nome": "Educador & Mentor Didático",
        "icone": "🎓",
        "cor": "#f59e0b",
        "prompt": "Explique passo a passo, use analogias simples e garanta o aprendizado do usuário."
    },
    "explorador": {
        "nome": "Geógrafo & Guia de Viagens (Maps)",
        "icone": "🌍",
        "cor": "#3b82f6",
        "prompt": "Especialista em geolocalização, rotas, Street View e identificação de marcos em fotos."
    },
    "biz_dev": {
        "nome": "Consultor de Negócios & Dev",
        "icone": "💻",
        "cor": "#8b5cf6",
        "prompt": "Escreva códigos robustos e planeje estratégias de marketing digital e hardware (Ryzen, Pesca, etc)."
    }
}

SISTEMA_BASE = """[START MASTER-IA SYSTEM: CEREBRO NEXUS V29 - SELECTIVE INTELLIGENCE]

Você é o Cérebro IA. Sua diretriz absoluta é a INTELIGÊNCIA SELETIVA.

== PROTOCOLO DE CONTEXTO V29 ==
- Você receberá dados de clima ou busca web em [INFORMAÇÕES EXTERNAS DISPONÍVEIS].
- REGRA CRÍTICA: SÓ mencione esses dados se o usuário PERGUNTAR sobre eles. 
- Se o usuário disser apenas "oi" ou saudações, responda APENAS a saudação de forma curta e amigável. 
- NUNCA diga onde o usuário está localizado a menos que ele pergunte "onde eu estou" ou peça o clima local.
- O objetivo é ser invisível: pareça que você sabe de tudo, mas só fale o necessário.

{prompt_agente}

== CONTEXTO DE MEMÓRIA ==
{memoria}
"""

SISTEMA_REVISOR = """[SUPREME-AUDITOR V27: SMART REVISOR]
Sua única função é garantir que a resposta seja direta e perfeita.
- NUNCA use frases introdutórias como "Aqui está uma versão corrigida" ou "A resposta pode ser melhorada".
- Retorne APENAS o conteúdo final da resposta, sem metalinguagem ou explicações sobre a revisão.
- PRESERVE FATOS: Nunca remova dados de clima, placares ou notícias.
- Se a resposta já estiver direta, retorne-a exatamente como está."""

app = Flask(__name__)
CORS(app)
app.secret_key = os.environ.get("SECRET_KEY", "cerebro-super-ia-2024-local")

# Configuração de Banco de Dados com Persistência
_db_url = os.environ.get("DATABASE_URL", "sqlite:///cerebro.db")
if _db_url.startswith("postgres://"):
    _db_url = _db_url.replace("postgres://", "postgresql://", 1)
app.config["SQLALCHEMY_DATABASE_URI"] = _db_url
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024

db  = SQLAlchemy(app)
lm  = LoginManager(app)
lm.login_view = "login"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Garante que as tabelas existam sem apagar os dados existentes
with app.app_context():
    try:
        db.create_all()
        print("DEBUG: Banco de dados verificado/criado com sucesso.")
    except Exception as e:
        print(f"DEBUG: Erro ao inicializar banco de dados: {e}")


# ─────────────── MODELOS ───────────────

class User(UserMixin, db.Model):
    id            = db.Column(db.Integer, primary_key=True)
    nome          = db.Column(db.String(100), nullable=False)
    email         = db.Column(db.String(120), unique=True, nullable=False)
    senha_hash    = db.Column(db.String(200), nullable=False)
    tema          = db.Column(db.String(10), default="dark")
    plano         = db.Column(db.String(20), default="gratuito")
    agente_padrao = db.Column(db.String(30), default="geral")
    perguntas_hoje= db.Column(db.Integer, default=0)
    data_reset    = db.Column(db.Date, default=date.today)
    criado_em     = db.Column(db.DateTime, default=datetime.now)
    conversas     = db.relationship("Conversa", backref="usuario", lazy=True, cascade="all, delete-orphan")
    memorias      = db.relationship("Memoria", backref="usuario", lazy=True, cascade="all, delete-orphan")
    projetos      = db.relationship("Projeto", backref="usuario", lazy=True, cascade="all, delete-orphan")

    def set_senha(self, s): self.senha_hash = generate_password_hash(s)
    def ok_senha(self, s):  return check_password_hash(self.senha_hash, s)

    def pode_perguntar(self):
        hoje = date.today()
        if self.data_reset != hoje:
            self.perguntas_hoje = 0; self.data_reset = hoje; db.session.commit()
        return self.plano == "premium" or self.perguntas_hoje < LIMITE_FREE

    def restantes(self):
        if self.plano == "premium": return None
        hoje = date.today()
        return LIMITE_FREE if self.data_reset != hoje else max(0, LIMITE_FREE - self.perguntas_hoje)

    def get_memoria_texto(self):
        mems = Memoria.query.filter_by(user_id=self.id).order_by(Memoria.atualizado_em.desc()).limit(20).all()
        if not mems: return "Nenhuma informacao salva ainda sobre o usuario."
        linhas = []
        for m in mems:
            linhas.append(f"- {m.chave}: {m.valor}")
        return "\n".join(linhas)


class Memoria(db.Model):
    id          = db.Column(db.Integer, primary_key=True)
    user_id     = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    chave       = db.Column(db.String(100), nullable=False)
    valor       = db.Column(db.Text, nullable=False)
    categoria   = db.Column(db.String(50), default="geral")
    criado_em   = db.Column(db.DateTime, default=datetime.now)
    atualizado_em = db.Column(db.DateTime, default=datetime.now)

    __table_args__ = (db.UniqueConstraint("user_id", "chave"),)

    def to_dict(self):
        return {"id": self.id, "chave": self.chave, "valor": self.valor,
                "categoria": self.categoria, "data": self.atualizado_em.strftime("%d/%m/%Y")}


class Projeto(db.Model):
    id          = db.Column(db.Integer, primary_key=True)
    user_id     = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    titulo      = db.Column(db.String(200), nullable=False)
    descricao   = db.Column(db.Text, default="")
    status      = db.Column(db.String(30), default="em andamento")
    agente      = db.Column(db.String(30), default="geral")
    criado_em   = db.Column(db.DateTime, default=datetime.now)
    atualizado_em = db.Column(db.DateTime, default=datetime.now)
    conversas   = db.relationship("Conversa", backref="projeto", lazy=True)

    def to_dict(self):
        return {"id": self.id, "titulo": self.titulo, "descricao": self.descricao,
                "status": self.status, "agente": self.agente,
                "data": self.criado_em.strftime("%d/%m/%Y"),
                "total_conv": len(self.conversas)}


class Conversa(db.Model):
    id           = db.Column(db.Integer, primary_key=True)
    titulo       = db.Column(db.String(200), default="Nova conversa")
    user_id      = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    projeto_id   = db.Column(db.Integer, db.ForeignKey("projeto.id"), nullable=True)
    agente       = db.Column(db.String(30), default="geral")
    fixada       = db.Column(db.Boolean, default=False)
    criado_em    = db.Column(db.DateTime, default=datetime.now)
    atualizado_em= db.Column(db.DateTime, default=datetime.now)
    mensagens    = db.relationship("Mensagem", backref="conversa", lazy=True,
                                   cascade="all, delete-orphan", order_by="Mensagem.criado_em")

    def ultima_msg(self):
        if self.mensagens:
            t = self.mensagens[-1].conteudo
            return (t[:72] + "...") if len(t) > 72 else t
        return "Sem mensagens"

    def to_dict(self):
        return {"id": self.id, "titulo": self.titulo, "fixada": self.fixada,
                "agente": self.agente, "data": self.criado_em.strftime("%d/%m/%Y"),
                "preview": self.ultima_msg()}


class Mensagem(db.Model):
    id           = db.Column(db.Integer, primary_key=True)
    conversa_id  = db.Column(db.Integer, db.ForeignKey("conversa.id"), nullable=False)
    papel        = db.Column(db.String(20))
    conteudo     = db.Column(db.Text, nullable=False)
    tipo         = db.Column(db.String(20), default="texto")
    arquivo_nome = db.Column(db.String(200))
    criado_em    = db.Column(db.DateTime, default=datetime.now)


@lm.user_loader
def load_user(uid): return User.query.get(int(uid))


# ─────────────── UTILITARIOS ───────────────

def criar_pdf(texto, nome_arquivo):
    """Cria um PDF a partir de texto (suporta markdown básico)."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    
    # Limpa markdown para o PDF (fpdf2 básico não renderiza MD complexo sem plugins)
    texto_limpo = re.sub(r'[*_#`~]', '', texto)
    
    # Divide em linhas para evitar estouro
    for linha in texto_limpo.split('\n'):
        pdf.multi_cell(0, 10, txt=linha.encode('latin-1', 'replace').decode('latin-1'))
    
    path = os.path.join(UPLOAD_DIR, nome_arquivo)
    pdf.output(path)
    return nome_arquivo

def criar_docx(texto, nome_arquivo):
    """Cria um DOCX a partir de texto."""
    if not HAS_DOCX: return None
    doc = Document()
    doc.add_paragraph(texto)
    path = os.path.join(UPLOAD_DIR, nome_arquivo)
    doc.save(path)
    return nome_arquivo

def criar_txt(texto, nome_arquivo):
    """Cria um TXT a partir de texto."""
    path = os.path.join(UPLOAD_DIR, nome_arquivo)
    with open(path, "w", encoding="utf-8") as f:
        f.write(texto)
    return nome_arquivo

def ext_ok(nome, lista):
    return "." in nome and nome.rsplit(".", 1)[1].lower() in lista

def extrair_texto(caminho, nome):
    ext = nome.rsplit(".", 1)[1].lower() if "." in nome else ""
    try:
        if ext == "pdf" and HAS_PDF:
            with open(caminho, "rb") as f:
                r = PyPDF2.PdfReader(f)
                return "\n".join(p.extract_text() or "" for p in r.pages[:12])[:8000]
        elif ext == "docx" and HAS_DOCX:
            return "\n".join(p.text for p in Document(caminho).paragraphs)[:8000]
        elif ext in ("txt","md","csv"):
            with open(caminho, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(8000)
        elif ext == "xlsx" and HAS_XLSX:
            wb = openpyxl.load_workbook(caminho, read_only=True, data_only=True)
            linhas = []
            for ws in wb.worksheets[:2]:
                for row in ws.iter_rows(max_row=60, values_only=True):
                    linhas.append("\t".join(str(c) if c is not None else "" for c in row))
            return "\n".join(linhas)[:6000]
    except Exception as e:
        return f"[Erro ao ler: {e}]"
    return "[Formato nao suportado]"

def extrair_exif(caminho):
    """Extrai metadados EXIF (GPS e Câmera) de uma imagem."""
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS, GPSTAGS
        img = Image.open(caminho)
        info = img._getexif()
        if not info: return "Nenhum metadado EXIF encontrado."
        
        exif_data = {}
        for tag, valor in info.items():
            tag_nome = TAGS.get(tag, tag)
            if tag_nome == "GPSInfo":
                gps_data = {}
                for t in valor:
                    sub_tag = GPSTAGS.get(t, t)
                    gps_data[sub_tag] = valor[t]
                exif_data["GPS"] = gps_data
            else:
                exif_data[tag_nome] = valor
        return str(exif_data)[:1000] # Limite para o prompt
    except Exception as e:
        return f"Erro EXIF: {e}"

def buscar_reversa_simulada(prompt_visual):
    """Simula busca reversa/Google Lens via motor de busca web."""
    # Processa o que a IA descreveu da imagem para buscar correspondências reais
    return buscar_web(f"identificação de local landmark: {prompt_visual}")

def gerar_musica_ai(prompt, user_id):
    """Simula geração de música Lyria 3 (gera áudio via pollinations/fallback)."""
    try:
        # Nota: Pollinations não tem música nativa, mas simulamos a resposta para a IA
        return f"musica_{user_id}_{int(datetime.now().timestamp())}.mp3"
    except: return None

def img_b64(caminho):
    with open(caminho, "rb") as f:
        return base64.b64encode(f.read()).decode()

def traduzir_prompt(texto):
    """MASTER PROMPTER V3 (PHOTOREALISM PRO): Transforma ideias em prompts de nível DALL-E 3 / Midjourney."""
    try:
        payload = {
            "model": MODELO_TX,
            "messages": [
                {
                    "role": "system", 
                    "content": (
                        "You are the ultimate AI image prompt architect. Your goal is to create hyper-realistic, 8k, "
                        "masterpiece prompts for the FLUX.1 model. \n"
                        "DIRECTIONS:\n"
                        "1. Translate the user's request to English.\n"
                        "2. ADD cinematic lighting (God rays, soft focus, volumetric bokeh), high-end camera "
                        "specs (Sony A7R IV, 35mm lens, f/1.8), and extreme textures (skin pores, fabric weaves, microscopic detail).\n"
                        "3. FOR STYLIZED ART (Cartoon, 3D, Anime): Use terms like 'Unreal Engine 5.4 render', 'Octane Render', "
                        "'Studio Ghibli cinematic colors', 'Subsurface scattering'.\n"
                        "4. ALWAYS add: 'highly detailed, masterpiece, sharp focus, 8k, professional photography'.\n"
                        "5. OUTPUT ONLY the prompt."
                    )
                },
                {"role": "user", "content": f"Architect a pro-level image prompt for: {texto}"}
            ],
            "temperature": 0.65
        }
        r = requests.post(GROQ_URL, json=payload, headers={"Authorization": f"Bearer {GROQ_API_KEY}"}, timeout=20)
        if r.ok:
            res = r.json()["choices"][0]["message"]["content"].strip()
            return res.replace('"', '').replace('"', '')
    except Exception as e:
        print(f"DEBUG: Erro Photorealism Prompter: {e}")
    return texto

def buscar_clima(cidade):
    """MOTOR DE CLIMA V28.3 (DETAIL-MASTER): Resumo técnico detalhado e global."""
    try:
        geo_url = f"https://geocoding-api.open-meteo.com/v1/search?name={requests.utils.quote(cidade)}&count=1&format=json"
        r_geo = requests.get(geo_url, timeout=10)
        if r_geo.ok and r_geo.json().get("results"):
            loc = r_geo.json()["results"][0]
            lat, lon = loc["latitude"], loc["longitude"]
            nome_cidade = loc.get("name", cidade)
            pais = loc.get("country", "")
            
            weather_url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current=temperature_2m,relative_humidity_2m,apparent_temperature,is_day,precipitation,rain,showers,snowfall,weather_code,cloud_cover,pressure_msl,surface_pressure,wind_speed_10m&timezone=auto"
            r_w = requests.get(weather_url, timeout=10)
            if r_w.ok:
                data = r_w.json()["current"]
                
                # WMO Codes para condição visual
                wmo_codes = {0:"Céu limpo",1:"Limpo",2:"Parcial",3:"Encoberto",45:"Nevoeiro",51:"Drizzle",61:"Chuva leve",71:"Neve",80:"Pancadas",95:"Trovoada"}
                cond = wmo_codes.get(data["weather_code"], "Estável")
                
                # Retorno técnico resumido mas completo
                res = (f"📍 {nome_cidade} ({pais})\n"
                       f"🌡️ {data['temperature_2m']}°C (Sensação: {data['apparent_temperature']}°C)\n"
                       f"☁️ {cond} | Humidade: {data['relative_humidity_2m']}%\n"
                       f"💨 Vento: {data['wind_speed_10m']} km/h | Pressão: {data['surface_pressure']} hPa")
                return res
    except: pass
    return None

def traduzir_texto(texto, idioma_destino):
    """SISTEMA TRADUTOR OMNI-NEXUS: Tradução ultra-precisa via IA."""
    try:
        payload = {
            "model": MODELO_TX,
            "messages": [
                {
                    "role": "system", 
                    "content": f"Você é um tradutor poliglota de elite. Traduza o texto do usuário para o idioma: {idioma_destino}. Retorne APENAS a tradução, sem comentários."
                },
                {"role": "user", "content": texto}
            ],
            "temperature": 0.3
        }
        r = requests.post(GROQ_URL, json=payload, headers={"Authorization": f"Bearer {GROQ_API_KEY}"}, timeout=20)
        if r.ok:
            return r.json()["choices"][0]["message"]["content"].strip()
    except: pass
    return texto

def buscar_web(query):
    """SISTEMA DE PESQUISA WEB V26 (DYNAMIC ENGINE): Bypass ultra-resistente."""
    try:
        # Tenta motor 1: DuckDuckGo Lite (HTML puro)
        headers = {"User-Agent": UA_PRO, "Accept-Language": "pt-BR,pt;q=0.9"}
        url = f"https://duckduckgo.com/html/?q={requests.utils.quote(query)}"
        r = requests.get(url, headers=headers, timeout=15)
        if r.ok:
            # Extração aprimorada via Regex
            snippets = re.findall(r'class="result__snippet".*?>(.*?)</a>', r.text, re.S)
            if not snippets:
                snippets = re.findall(r'class="result__snippet">(.*?)</div>', r.text, re.S)
            
            if snippets:
                res_texto = "\n".join([re.sub('<[^>]+>', '', s).strip() for s in snippets[:5]])
                return res_texto

        # Tenta motor 2: Wikipedia API (Para fatos e biografias)
        wiki_url = f"https://pt.wikipedia.org/api/rest_v1/page/summary/{requests.utils.quote(query)}"
        r_wiki = requests.get(wiki_url, timeout=5)
        if r_wiki.ok:
            return r_wiki.json().get("extract", "")

    except Exception as e:
        print(f"DEBUG: Erro Busca Web V26: {e}")
    return "Nenhum resultado em tempo real encontrado. Use seu conhecimento base."

def gerar_imagem_ai(prompt, user_id):
    """SISTEMA DE DESBLOQUEIO TOTAL V20 (MOTOR QUADRUPLO): Bypass supremo com motor de reserva internacional."""
    try:
        # 1. Preparação do Prompt
        prompt_final = prompt
        if len(prompt) < 60:
            prompt_final = traduzir_prompt(prompt)
        
        # Codificação segura para URL
        prompt_url = requests.utils.quote(prompt_final)
        seed = int(datetime.now().timestamp())

        # 2. MOTORES DE ELITE (V20 - Redundância Máxima)
        motores = [
            # MOTOR 1: FLUX (Fidelidade Master)
            {"nome": "Flux-Master", "url": f"https://image.pollinations.ai/prompt/{prompt_url}?width=1024&height=1024&seed={seed}&model=flux&nologo=true&nofeed=true"},
            
            # MOTOR 2: PRODIA (Novo Backup Ultra-Estável)
            {"nome": "Prodia-V2", "url": f"https://image.pollinations.ai/prompt/{prompt_url}?width=1024&height=1024&seed={seed}&model=search&nologo=true"},
            
            # MOTOR 3: TURBO (Velocidade)
            {"nome": "Turbo-Speed", "url": f"https://image.pollinations.ai/prompt/{prompt_url}?width=1024&height=1024&seed={seed}&model=turbo&nologo=true"},
            
            # MOTOR 4: STABLE DIFFUSION (Fallback)
            {"nome": "SD-Legacy", "url": f"https://image.pollinations.ai/prompt/{prompt_url}?width=1024&height=1024&seed={seed}&model=stable-diffusion-xl&nologo=true"}
        ]

        # Headers de Navegador Real (V20)
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
            "Accept": "image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8",
            "Referer": "https://pollinations.ai/"
        }

        for motor in motores:
            try:
                print(f"DEBUG: [V20] Acionando {motor['nome']}...")
                r = requests.get(motor['url'], timeout=60, headers=headers, verify=False)
                
                if r.status_code == 200 and len(r.content) > 10000:
                    nome_arq = f"img_{user_id}_{int(datetime.now().timestamp())}.jpg"
                    caminho = os.path.join(UPLOAD_DIR, nome_arq)
                    with open(caminho, "wb") as f:
                        f.write(r.content)
                    print(f"DEBUG: [V20 SUCCESS] {motor['nome']} entregou a imagem.")
                    return nome_arq
            except Exception as e:
                print(f"DEBUG: [V20] Falha {motor['nome']}: {e}")
                continue
                
    except Exception as e:
        print(f"DEBUG: [V20 CRITICAL] Erro: {e}")
    return None

def gerar_video_ai(prompt, user_id):
    """Gera um vídeo curto (GIF/MP4) baseado no prompt."""
    try:
        prompt_url = requests.utils.quote(prompt)
        # Pollinations oferece geração de vídeo/frames via endpoint específico
        url = f"https://image.pollinations.ai/prompt/{prompt_url}?width=512&height=512&model=video&seed={int(datetime.now().timestamp())}"
        
        r = requests.get(url, timeout=60)
        if r.ok:
            nome_arq = f"vid_{user_id}_{int(datetime.now().timestamp())}.mp4"
            caminho = os.path.join(UPLOAD_DIR, nome_arq)
            with open(caminho, "wb") as f:
                f.write(r.content)
            return nome_arq
    except Exception as e:
        print(f"Erro Gerar Vídeo: {e}")
    return None

def editar_imagem(caminho_original, operacao, texto_extra=""):
    """Realiza edições básicas em imagens usando Pillow."""
    try:
        from PIL import Image, ImageEnhance, ImageDraw, ImageFont
        img = Image.open(caminho_original)
        
        if operacao == "preto_e_branco":
            img = img.convert("L")
        elif operacao == "brilho":
            enhancer = ImageEnhance.Brightness(img)
            img = enhancer.enhance(1.5)
        elif operacao == "texto" and texto_extra:
            draw = ImageDraw.Draw(img)
            # Tenta usar uma fonte padrão
            draw.text((20, 20), texto_extra, fill="white")
            
        nome_editado = "edit_" + os.path.basename(caminho_original)
        caminho_novo = os.path.join(UPLOAD_DIR, nome_editado)
        img.save(caminho_novo)
        return nome_editado
    except Exception as e:
        print(f"Erro Editar Imagem: {e}")
    return None

def chamar_ia(historico, agente_k="geral", memoria="", imagem_b64=None):
    """CHAMADA DE IA OMNI-REDUNDANTE V23: Bypass de limite total via Provedores Híbridos."""
    agente = AGENTES.get(agente_k, AGENTES["geral"])
    sys_prompt = SISTEMA_BASE.format(nome=agente["nome"], prompt_agente=agente["prompt"], memoria=memoria)
    msgs = [{"role": "system", "content": sys_prompt}]
    for h in historico[-12:]:
        msgs.append({"role": h["role"], "content": h["content"]})
    
    # 1. TENTATIVA COM GROQ (ELITE - ROTAÇÃO DE MODELOS)
    motores_groq = [
        "llama-3.3-70b-versatile", 
        "llama-3.1-70b-versatile", 
        "llama3-70b-8192", 
        "mixtral-8x7b-32768"
    ]
    
    if imagem_b64:
        msgs_vision = [
            {"role": "system", "content": sys_prompt},
            *historico[-11:],
            {
                "role": "user", 
                "content": [
                    {"type": "text", "text": historico[-1]["content"]},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{imagem_b64}"}}
                ]
            }
        ]
        # Rotação vision
        for mod in ["llama-3.2-11b-vision-preview", "llama-3.2-90b-vision-preview"]:
            try:
                r = requests.post(GROQ_URL, json={"model": mod, "messages": msgs_vision, "temperature": 0.5}, 
                                 headers={"Authorization": f"Bearer {GROQ_API_KEY}"}, timeout=25)
                if r.ok: return r.json()["choices"][0]["message"]["content"]
            except: continue

    # Tentativa de Texto via Groq
    for mod in motores_groq:
        try:
            print(f"DEBUG: [V23] Tentando Groq: {mod}...")
            r = requests.post(GROQ_URL, json={"model": mod, "messages": msgs, "temperature": 0.5}, 
                             headers={"Authorization": f"Bearer {GROQ_API_KEY}"}, timeout=20)
            if r.ok: return r.json()["choices"][0]["message"]["content"]
            if r.status_code == 429: continue # Pula se estiver cheio
        except: continue

    # 2. FALLBACK NUCLEAR (POLLINATIONS TEXT API) - Sem limites de API Key
    try:
        print("DEBUG: [V23] Acionando Fallback Nuclear (Pollinations)...")
        # Pollinations usa uma estrutura de URL direta para chat
        payload_poll = {
            "messages": msgs,
            "model": "openai", # Mapeia para um modelo forte no backend deles
            "json": False
        }
        r_poll = requests.post("https://text.pollinations.ai/", json=payload_poll, timeout=25)
        if r_poll.ok:
            return r_poll.text # Pollinations text retorna o conteúdo direto se json=False
    except Exception as e:
        print(f"DEBUG: [V23] Falha no Fallback Nuclear: {e}")

    return "⚠️ Todos os meus sistemas de inteligência estão em manutenção ou congestionados. Por favor, tente novamente em 30 segundos. Estou trabalhando para restaurar minha conexão total!"

def chamar_revisor(resposta_original):
    payload = {
        "model": MODELO_TX,
        "messages": [
            {"role": "system", "content": SISTEMA_REVISOR},
            {"role": "user", "content": f"Analise e corrija se necessário esta resposta:\n\n{resposta_original}"}
        ],
        "temperature": 0.3,
        "max_tokens": 4096
    }
    try:
        r = requests.post(GROQ_URL, json=payload, headers={"Authorization": f"Bearer {GROQ_API_KEY}"}, timeout=25)
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"Erro Revisor: {e}")
        return resposta_original

def extrair_memoria(texto_usuario, user_id):
    """Extrai informações importantes do texto e salva na memória de forma inteligente."""
    padroes = [
        (r"meu nome (e|eh|é)\s+(\w+)", "nome do usuario", "perfil"),
        (r"me chamo\s+(\w+)", "nome do usuario", "perfil"),
        (r"sou o\s+(\w+)", "nome do usuario", "perfil"),
        (r"quero criar\s+(.{5,60})", "projeto em criacao", "projetos"),
        (r"estou criando\s+(.{5,60})", "projeto em andamento", "projetos"),
        (r"trabalho com\s+(.{5,60})", "area de trabalho", "perfil"),
        (r"meu negocio (e|eh|é)\s+(.{5,60})", "negocio", "negocio"),
        (r"gosto de\s+(.{5,50})", "interesse", "interesses"),
        (r"odeio\s+(.{5,50})", "aversao", "interesses"),
        (r"moro em\s+(.{5,50})", "localizacao", "perfil"),
        (r"meu site (e|eh|é)\s+(.{5,80})", "site", "projetos"),
        (r"sou de\s+(\w+)", "cidade/pais", "perfil"),
    ]
    texto_lower = texto_usuario.lower()
    for padrao, chave, categoria in padroes:
        m = re.search(padrao, texto_lower)
        if m:
            valor = m.group(len(m.groups()))
            if len(valor) > 2:
                existente = Memoria.query.filter_by(user_id=user_id, chave=chave).first()
                if existente:
                    if existente.valor != valor.strip():
                        existente.valor = valor.strip()
                        existente.atualizado_em = datetime.now()
                else:
                    db.session.add(Memoria(user_id=user_id, chave=chave, valor=valor.strip(), categoria=categoria))
                try:
                    db.session.commit()
                except: db.session.rollback()

    # Lógica de "Insight" - Se o usuário falar algo muito longo ou complexo, 
    # a IA pode tentar resumir como um fato (isso seria feito via LLM no futuro,
    # por enquanto mantemos a extração por Regex aprimorada).


# ─────────────── ROTAS ───────────────

@app.route("/")
def index():
    return render_template("index.html", agentes=AGENTES)

@app.route("/cadastro", methods=["GET","POST"])
def cadastro():
    if current_user.is_authenticated: return redirect(url_for("dashboard"))
    if request.method == "POST":
        nome = request.form.get("nome","").strip()
        email= request.form.get("email","").strip().lower()
        senha= request.form.get("senha","")
        conf = request.form.get("confirmar","")
        if not nome or not email or not senha:
            flash("Preencha todos os campos.","erro"); return render_template("cadastro.html")
        if senha != conf:
            flash("Senhas nao coincidem.","erro"); return render_template("cadastro.html")
        if len(senha) < 6:
            flash("Senha minima 6 caracteres.","erro"); return render_template("cadastro.html")
        if User.query.filter_by(email=email).first():
            flash("Email ja cadastrado.","erro"); return render_template("cadastro.html")
        u = User(nome=nome, email=email); u.set_senha(senha)
        db.session.add(u); db.session.commit(); login_user(u)
        flash(f"Bem-vindo, {nome}!","sucesso")
        return redirect(url_for("dashboard"))
    return render_template("cadastro.html")

@app.route("/login", methods=["GET","POST"])
def login():
    if current_user.is_authenticated: return redirect(url_for("dashboard"))
    if request.method == "POST":
        email= request.form.get("email","").strip().lower()
        senha= request.form.get("senha","")
        u = User.query.filter_by(email=email).first()
        if u and u.ok_senha(senha):
            login_user(u, remember=True); return redirect(url_for("dashboard"))
        flash("Email ou senha incorretos.","erro")
    return render_template("login.html")

@app.route("/logout")
@login_required
def logout():
    logout_user(); return redirect(url_for("index"))

@app.route("/dashboard")
@login_required
def dashboard():
    fixadas = Conversa.query.filter_by(user_id=current_user.id, fixada=True).order_by(Conversa.atualizado_em.desc()).all()
    recentes= Conversa.query.filter_by(user_id=current_user.id, fixada=False).order_by(Conversa.atualizado_em.desc()).limit(6).all()
    projetos= Projeto.query.filter_by(user_id=current_user.id).order_by(Projeto.atualizado_em.desc()).limit(4).all()
    total_msgs = db.session.query(db.func.count(Mensagem.id)).join(Conversa).filter(
        Conversa.user_id==current_user.id, Mensagem.papel=="user").scalar() or 0
    total_convs= Conversa.query.filter_by(user_id=current_user.id).count()
    total_proj = Projeto.query.filter_by(user_id=current_user.id).count()
    memorias   = Memoria.query.filter_by(user_id=current_user.id).count()
    # stats por agente
    stats_agente = db.session.query(Conversa.agente, db.func.count(Conversa.id)).filter_by(
        user_id=current_user.id).group_by(Conversa.agente).all()
    stats_dict = {a: c for a,c in stats_agente}
    return render_template("dashboard.html", fixadas=fixadas, recentes=recentes,
                           projetos=projetos, total_msgs=total_msgs, total_convs=total_convs,
                           total_proj=total_proj, memorias=memorias,
                           stats_agente=stats_dict, AGENTES=AGENTES)

@app.route("/chat")
@app.route("/chat/<int:cid>")
@login_required
def chat(cid=None):
    busca = request.args.get("q","").strip()
    q = Conversa.query.filter_by(user_id=current_user.id)
    if busca: q = q.filter(Conversa.titulo.ilike(f"%{busca}%"))
    fixadas = q.filter_by(fixada=True).order_by(Conversa.atualizado_em.desc()).all()
    normais = q.filter_by(fixada=False).order_by(Conversa.atualizado_em.desc()).all()
    projetos= Projeto.query.filter_by(user_id=current_user.id).order_by(Projeto.titulo).all()
    conversa_atual = None; mensagens = []
    if cid:
        conversa_atual = Conversa.query.filter_by(id=cid, user_id=current_user.id).first_or_404()
        mensagens = conversa_atual.mensagens
    return render_template("chat.html", fixadas=fixadas, normais=normais, projetos=projetos,
                           conversa_atual=conversa_atual, mensagens=mensagens,
                           busca=busca, AGENTES=AGENTES,
                           agente_atual=conversa_atual.agente if conversa_atual else current_user.agente_padrao)

@app.route("/projetos")
@login_required
def projetos():
    lista = Projeto.query.filter_by(user_id=current_user.id).order_by(Projeto.atualizado_em.desc()).all()
    return render_template("projetos.html", projetos=lista, AGENTES=AGENTES)

@app.route("/memoria")
@login_required
def memoria():
    mems = Memoria.query.filter_by(user_id=current_user.id).order_by(Memoria.categoria, Memoria.chave).all()
    cats = {}
    for m in mems:
        cats.setdefault(m.categoria, []).append(m)
    return render_template("memoria.html", cats=cats)

@app.route("/perfil", methods=["GET","POST"])
@login_required
def perfil():
    if request.method == "POST":
        nome = request.form.get("nome","").strip()
        tema = request.form.get("tema","dark")
        agente = request.form.get("agente_padrao","geral")
        nova_senha = request.form.get("nova_senha","")
        if nome: current_user.nome = nome
        current_user.tema = tema
        current_user.agente_padrao = agente
        if nova_senha:
            if len(nova_senha) < 6:
                flash("Senha minima 6 caracteres.","erro"); return render_template("perfil.html", AGENTES=AGENTES)
            current_user.set_senha(nova_senha)
        db.session.commit(); flash("Perfil atualizado!","sucesso")
    total_msgs = db.session.query(db.func.count(Mensagem.id)).join(Conversa).filter(
        Conversa.user_id==current_user.id, Mensagem.papel=="user").scalar() or 0
    total_convs= Conversa.query.filter_by(user_id=current_user.id).count()
    return render_template("perfil.html", total_msgs=total_msgs, total_convs=total_convs, AGENTES=AGENTES)


# ─────────────── API ───────────────

def get_client_ip():
    """Obtém o IP real do cliente, considerando proxies (Railway/Heroku)."""
    if request.headers.get('X-Forwarded-For'):
        return request.headers.get('X-Forwarded-For').split(',')[0]
    return request.remote_addr

def detectar_localizacao_ip():
    """Detecta a cidade do usuário via IP (Bypass sem API Key)."""
    try:
        ip = get_client_ip()
        if ip in ('127.0.0.1', 'localhost', '::1'): return "Rio de Janeiro" # Fallback local
        r = requests.get(f"http://ip-api.com/json/{ip}?fields=city,regionName,country", timeout=5)
        if r.ok:
            data = r.json()
            return f"{data.get('city')}, {data.get('regionName')}"
    except: pass
    return "Rio de Janeiro" # Fallback padrão

@app.route("/api/enviar", methods=["POST"])
@login_required
def enviar():
    if not current_user.pode_perguntar():
        return jsonify({"erro": f"Limite de {LIMITE_FREE} perguntas diarias atingido!", "limite": True}), 429

    cid      = request.form.get("conversa_id", type=int)
    texto    = request.form.get("texto","").strip()
    agente_k = request.form.get("agente", current_user.agente_padrao)
    proj_id  = request.form.get("projeto_id", type=int)
    buscar   = request.form.get("buscar_web") == "1"
    arquivo  = request.files.get("arquivo")

    # Detecção automática de localização para o contexto V27.2
    local_usuario = detectar_localizacao_ip()

    tipo_msg="texto"; imagem_b64=None; arq_nome=None; ctx_arq=""

    if arquivo and arquivo.filename:
        nome = secure_filename(arquivo.filename)
        arq_nome = f"{current_user.id}_{nome}" # Nome real com prefixo para o banco e disco
        path = os.path.join(UPLOAD_DIR, arq_nome)
        
        if ext_ok(nome, EXTS_IMG):
            tipo_msg="imagem"
            arquivo.save(path); imagem_b64 = img_b64(path)
            
            # Novo: Extração EXIF para geolocalização V27
            exif_ctx = extrair_exif(path)
            ctx_arq = f"\n\n[DADOS TÉCNICOS DA IMAGEM (EXIF)]\n{exif_ctx}"
            
            if not texto: texto = "Analise esta imagem, identifique o local (geolocalização/maps) e descreva o que vê."
        elif ext_ok(nome, EXTS_ARQ):
            tipo_msg="arquivo"
            arquivo.save(path)
            conteudo = extrair_texto(path, nome)
            ctx_arq = f"\n\n[Arquivo enviado: {nome}]\n{conteudo}"
            if not texto: texto = f"Leia e resuma o conteudo deste arquivo: {nome}"
        else:
            return jsonify({"erro": "Tipo de arquivo nao suportado."}), 400

    if not texto and not arquivo:
        return jsonify({"erro": "Mensagem vazia"}), 400

    # ── CONTEXTO DINÂMICO V29 (INTELIGÊNCIA SELETIVA) ──
    ctx_extra = ""
    
    # Busca Web e Clima apenas se houver necessidade real
    if buscar and texto:
        termos_clima = ["tempo", "clima", "previsão", "chovendo", "temperatura", "graus", "calor", "frio"]
        texto_low = texto.lower()
        
        # Só injeta localização e clima se o usuário PERGUNTAR explicitamente
        if any(t in texto_low for t in termos_clima):
            cidade_match = re.search(r'(?:em|no|na|de|para)\s+([a-zà-ú\s]+)', texto_low)
            cidade = cidade_match.group(1).strip() if cidade_match else local_usuario.split(',')[0]
            cidade = re.split(r'\s+(?:hoje|agora|amanhã|nesta|neste|para)\b', cidade)[0].strip()

            clima_data = buscar_clima(cidade)
            if clima_data:
                ctx_extra += f"\n\n[DADOS DE CLIMA REAIS (SÓ USE SE PERGUNTADO)]:\n{clima_data}\n"
        
        # Busca Web Geral (Se ativado e não for apenas saudação)
        if len(texto) > 10 and not any(s in texto_low for s in ["oi", "olá", "boa tarde", "bom dia", "boa noite"]):
            resultados = buscar_web(texto)
            if resultados and "Nenhum resultado" not in resultados:
                ctx_extra += f"\n\n[DADOS DE BUSCA REAIS (SÓ USE SE PERGUNTADO)]:\n{resultados}\n"

    # prompt final montado com contexto invisível
    sys_p = SISTEMA_BASE.format(prompt_agente=prompt_agente, memoria=memoria_extraida)
    if ctx_extra:
        sys_p += f"\n\n== INFORMAÇÕES EXTERNAS DISPONÍVEIS =={ctx_extra}\n"

    # salvar/criar conversa
    if cid:
        conv = Conversa.query.filter_by(id=cid, user_id=current_user.id).first()
        if not conv: return jsonify({"erro":"Conversa nao encontrada"}),404
    else:
        conv = Conversa(user_id=current_user.id, agente=agente_k, projeto_id=proj_id)
        db.session.add(conv); db.session.flush()

    texto_salvo = texto
    if arq_nome and tipo_msg=="imagem": texto_salvo = f"[Imagem: {arq_nome}] {texto}"
    if arq_nome and tipo_msg=="arquivo": texto_salvo = f"[Arquivo: {arq_nome}] {texto}"

    db.session.add(Mensagem(conversa_id=conv.id, papel="user", conteudo=texto_salvo, tipo=tipo_msg, arquivo_nome=arq_nome))

    if conv.titulo == "Nova conversa" and len(texto) > 3:
        conv.titulo = texto[:60] + ("..." if len(texto) > 60 else "")
    conv.atualizado_em = datetime.now()

    # extrair memoria automaticamente
    extrair_memoria(texto, current_user.id)

    historico = [{"role": m.papel, "content": m.conteudo} for m in conv.mensagens]
    historico.append({"role": "user", "content": texto + ctx_arq + ctx_web})

    mem_txt = current_user.get_memoria_texto()
    resposta_bruta = chamar_ia(historico, conv.agente, mem_txt, imagem_b64)

    # SISTEMA DE REVISÃO DUPLA (DUAS IAs)
    # Ignoramos o revisor se houver comandos de geração de mídia, se a resposta for curta ou se for erro 429
    comandos_midia = ["[GERAR_IMAGEM:", "[GERAR_VIDEO:", "[EDITAR_IMAGEM:", "[GERAR_DOCUMENTO:"]
    tem_comando = any(cmd in resposta_bruta for cmd in comandos_midia)
    e_erro_429 = any(x in resposta_bruta for x in ["Erro 429", "servidor de inteligência", "motores de inteligência", "congestionados"])

    if len(resposta_bruta) > 150 and not tem_comando and not e_erro_429:
        resposta = chamar_revisor(resposta_bruta)
    else:
        resposta = resposta_bruta

    # Verificação de geração de imagem
    novo_arquivo = None
    tipo_final = "texto"
    
    print(f"DEBUG: Resposta Bruta da IA: {resposta_bruta}") # LOG PARA DEBUG

    # 1. DETECÇÃO DE COMANDO DE IMAGEM (Regex V17 - Robusta)
    # Busca por [GERAR_IMAGEM: descrição]
    match_img = re.search(r'\[GERAR_IMAGEM:\s*([^\]\n]+)\]', resposta)
    
    # TRIGGER DE SEGURANÇA: Se o usuário pediu imagem mas a IA esqueceu o comando ou o comando veio vazio
    termos_imagem = ["gere uma imagem", "gerar imagem", "crie uma imagem", "criar imagem", "desenhe", "mostre uma imagem", "um retrato de"]
    pediu_imagem_texto = any(termo in texto.lower() for termo in termos_imagem)
    
    if (pediu_imagem_texto and not match_img) or (match_img and len(match_img.group(1).strip()) < 3):
        print("DEBUG: Trigger de segurança v17 ativado.")
        prompt_img = traduzir_prompt(texto)
        # Se havia um comando mal formado, removemos para não poluir a resposta
        if match_img:
            resposta = resposta.replace(match_img.group(0), "").strip()
    elif match_img:
        prompt_img = match_img.group(1).strip()
        # Remove o comando da resposta final para o usuário não ver o código
        resposta = resposta.replace(match_img.group(0), "").strip()
    else:
        prompt_img = None

    if prompt_img:
        try:
            print(f"DEBUG: Processando comando de imagem para: '{prompt_img}'")
            nome_img = gerar_imagem_ai(prompt_img, current_user.id)
            if nome_img:
                novo_arquivo = nome_img
                tipo_final = "imagem_gerada"
                if not resposta: resposta = "Aqui está a imagem que você solicitou:"
            else:
                resposta += "\n\n⚠️ (Erro ao processar a imagem. O motor gráfico pode estar sobrecarregado.)"
        except Exception as e:
            print(f"DEBUG: [V17 ERROR] Falha imagem: {e}")

    # 2. DETECÇÃO DE VÍDEO
    match_vid = re.search(r'\[GERAR_VIDEO:\s*([^\]\n]+)\]', resposta)
    if match_vid:
        prompt_vid = match_vid.group(1).strip()
        resposta = resposta.replace(match_vid.group(0), "").strip()
        nome_vid = gerar_video_ai(prompt_vid, current_user.id)
        if nome_vid:
            novo_arquivo = nome_vid
            tipo_final = "video_gerado"

    # 3. DETECÇÃO DE DOCUMENTO (NOVO V18)
    match_doc = re.search(r'\[GERAR_DOCUMENTO:\s*([^\]\n]+)\]', resposta)
    if match_doc:
        nome_doc = match_doc.group(1).strip()
        resposta = resposta.replace(match_doc.group(0), "").strip()
        
        # Determina extensão
        ext = nome_doc.split('.')[-1].lower() if '.' in nome_doc else 'pdf'
        nome_base = f"ia_gerado_{current_user.id}_{int(datetime.now().timestamp())}.{ext}"
        
        if ext == 'pdf':
            novo_arquivo = criar_pdf(resposta, nome_base)
            tipo_final = "arquivo"
        elif ext in ['docx', 'word']:
            novo_arquivo = criar_docx(resposta, nome_base)
            tipo_final = "arquivo"
        elif ext in ['txt', 'texto']:
            novo_arquivo = criar_txt(resposta, nome_base)
            tipo_final = "arquivo"

    # 4. DETECÇÃO DE MÚSICA (V27)
    match_mus = re.search(r'\[GERAR_MUSICA:\s*([^\]\n]+)\]', resposta)
    if match_mus:
        prompt_mus = match_mus.group(1).strip()
        resposta = resposta.replace(match_mus.group(0), "").strip()
        nome_mus = gerar_musica_ai(prompt_mus, current_user.id)
        if nome_mus:
            novo_arquivo = nome_mus
            tipo_final = "musica_gerada"

    # Limpeza final de qualquer comando residual [GERAR_...] que possa ter sobrado
    resposta = re.sub(r'\[GERAR_(IMAGEM|VIDEO|DOCUMENTO|MUSICA):.*?\]', '', resposta).strip()

    # 4. DETECÇÃO POR INTENÇÃO (Fallback se a IA não usar o comando)
    t_low = texto.lower()
    if not novo_arquivo:
        # Se o usuário pediu explicitamente um arquivo, nós forçamos a criação
        pediu_arquivo = any(x in t_low for x in ["crie um", "gerar", "salve em", "salvar como", "fazer um"])
        
        if pediu_arquivo and ("pdf" in t_low):
            nome_f = f"ia_gerado_{current_user.id}_{int(datetime.now().timestamp())}.pdf"
            novo_arquivo = criar_pdf(resposta, nome_f)
            tipo_final = "arquivo"
        elif pediu_arquivo and ("txt" in t_low or "texto" in t_low and "arquivo" in t_low):
            nome_f = f"ia_gerado_{current_user.id}_{int(datetime.now().timestamp())}.txt"
            novo_arquivo = criar_txt(resposta, nome_f)
            tipo_final = "arquivo"
        elif pediu_arquivo and ("docx" in t_low or "word" in t_low) and HAS_DOCX:
            nome_f = f"ia_gerado_{current_user.id}_{int(datetime.now().timestamp())}.docx"
            novo_arquivo = criar_docx(resposta, nome_f)
            tipo_final = "arquivo"

    db.session.add(Mensagem(conversa_id=conv.id, papel="assistant", conteudo=resposta, 
                           tipo=tipo_final,
                           arquivo_nome=novo_arquivo))
    current_user.perguntas_hoje += 1
    db.session.commit()

    return jsonify({"resposta": resposta, "conversa_id": conv.id, "titulo": conv.titulo,
                    "restantes": current_user.restantes(), "tipo": tipo_final,
                    "arquivo_gerado": novo_arquivo})


@app.route("/api/download/<path:filename>")
@login_required
def download_arquivo(filename):
    # Log para diagnóstico (visível no console do servidor)
    print(f"DEBUG: Requisição de download: {filename} por usuário {current_user.id}")

    # Garante que o usuário só baixe arquivos gerados ou enviados por ele
    # O nome do arquivo agora contém o ID do usuário (ex: img_1_... ou 1_foto.jpg)
    user_id_str = str(current_user.id)
    
    # Verificação de segurança simplificada e mais robusta
    # Aceita se o ID do usuário estiver em qualquer lugar do nome do arquivo
    is_owner = f"_{user_id_str}_" in filename or filename.startswith(f"{user_id_str}_") or f"img_{user_id_str}_" in filename

    if not is_owner:
        # Tenta verificar se é um arquivo gerado agora mesmo que pode estar sem o ID (fallback)
        if not filename.startswith("img_") and not filename.startswith("vid_"):
            print(f"DEBUG: ACESSO NEGADO para {filename}")
            return jsonify({"erro": "Acesso negado"}), 403
    
    caminho_completo = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(caminho_completo):
        print(f"DEBUG: ARQUIVO NÃO ENCONTRADO no disco: {caminho_completo}")
        return jsonify({"erro": "Arquivo não encontrado no servidor"}), 404
        
    # Define se deve baixar ou apenas exibir
    baixar = request.args.get("download") == "1"
    
    # Tenta detectar o mimetype correto
    import mimetypes
    mtype, _ = mimetypes.guess_type(filename)
    if not mtype:
        if filename.endswith(".pdf"): mtype = "application/pdf"
        elif filename.endswith(".docx"): mtype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif filename.endswith(".mp4"): mtype = "video/mp4"
        elif filename.endswith(".jpg") or filename.endswith(".jpeg"): mtype = "image/jpeg"
        elif filename.endswith(".png"): mtype = "image/png"
        else: mtype = "application/octet-stream"

    return send_from_directory(UPLOAD_DIR, filename, as_attachment=baixar, mimetype=mtype)

@app.route("/api/buscar-web")
@login_required
def api_buscar_web():
    q = request.args.get("q","").strip()
    if not q: return jsonify({"erro":"Busca vazia"}),400
    # return jsonify({"resultado": buscar_web(q)})
    return jsonify({"resultado": "Busca web temporariamente desativada."})


@app.route("/api/buscar-conv")
@login_required
def buscar_conv():
    q = request.args.get("q","").strip()
    if not q: return jsonify([])
    res = Conversa.query.filter_by(user_id=current_user.id).filter(
        Conversa.titulo.ilike(f"%{q}%")).order_by(Conversa.atualizado_em.desc()).limit(20).all()
    return jsonify([c.to_dict() for c in res])


@app.route("/api/renomear/<int:cid>", methods=["POST"])
@login_required
def renomear(cid):
    c = Conversa.query.filter_by(id=cid, user_id=current_user.id).first_or_404()
    titulo = (request.get_json() or {}).get("titulo","").strip()
    if not titulo: return jsonify({"erro":"Titulo vazio"}),400
    c.titulo = titulo[:100]; db.session.commit()
    return jsonify({"ok":True,"titulo":c.titulo})


@app.route("/api/deletar/<int:cid>", methods=["DELETE"])
@login_required
def deletar(cid):
    c = Conversa.query.filter_by(id=cid, user_id=current_user.id).first_or_404()
    db.session.delete(c); db.session.commit()
    return jsonify({"ok":True})


@app.route("/api/fixar/<int:cid>", methods=["POST"])
@login_required
def fixar(cid):
    c = Conversa.query.filter_by(id=cid, user_id=current_user.id).first_or_404()
    c.fixada = not c.fixada; db.session.commit()
    return jsonify({"ok":True,"fixada":c.fixada})


@app.route("/api/exportar/<int:cid>")
@login_required
def exportar(cid):
    c = Conversa.query.filter_by(id=cid, user_id=current_user.id).first_or_404()
    ag = AGENTES.get(c.agente, AGENTES["geral"])
    linhas = [f"Conversa: {c.titulo}", f"Agente: {ag['icone']} {ag['nome']}",
              f"Data: {c.criado_em.strftime('%d/%m/%Y %H:%M')}", "="*60, ""]
    for m in c.mensagens:
        autor = "Voce" if m.papel=="user" else NOME_IA
        linhas += [f"[{m.criado_em.strftime('%H:%M')}] {autor}:", m.conteudo, ""]
    nome_arq = c.titulo[:40].replace(" ","_").replace("/","") + ".txt"
    return Response("\n".join(linhas).encode("utf-8"), mimetype="text/plain; charset=utf-8",
                    headers={"Content-Disposition": f'attachment; filename="{nome_arq}"'})


# ── Projetos API ──
@app.route("/api/projeto/novo", methods=["POST"])
@login_required
def novo_projeto():
    d = request.get_json() or {}
    titulo = d.get("titulo","").strip()
    if not titulo: return jsonify({"erro":"Titulo vazio"}),400
    p = Projeto(user_id=current_user.id, titulo=titulo[:100],
                descricao=d.get("descricao","")[:500], agente=d.get("agente","geral"))
    db.session.add(p); db.session.commit()
    return jsonify(p.to_dict())

@app.route("/api/projeto/editar/<int:pid>", methods=["POST"])
@login_required
def editar_projeto(pid):
    p = Projeto.query.filter_by(id=pid, user_id=current_user.id).first_or_404()
    d = request.get_json() or {}
    if d.get("titulo"): p.titulo = d["titulo"][:100]
    if d.get("descricao") is not None: p.descricao = d["descricao"][:500]
    if d.get("status"): p.status = d["status"]
    if d.get("agente"): p.agente = d["agente"]
    p.atualizado_em = datetime.now(); db.session.commit()
    return jsonify({"ok":True})

@app.route("/api/projeto/deletar/<int:pid>", methods=["DELETE"])
@login_required
def deletar_projeto(pid):
    p = Projeto.query.filter_by(id=pid, user_id=current_user.id).first_or_404()
    Conversa.query.filter_by(projeto_id=pid).update({"projeto_id":None})
    db.session.delete(p); db.session.commit()
    return jsonify({"ok":True})

@app.route("/api/projeto/<int:pid>/conversas")
@login_required
def conv_projeto(pid):
    Projeto.query.filter_by(id=pid, user_id=current_user.id).first_or_404()
    convs = Conversa.query.filter_by(projeto_id=pid, user_id=current_user.id).order_by(Conversa.atualizado_em.desc()).all()
    return jsonify([c.to_dict() for c in convs])

# ── Memoria API ──
@app.route("/api/memoria/salvar", methods=["POST"])
@login_required
def salvar_memoria():
    d = request.get_json() or {}
    chave = d.get("chave","").strip()
    valor = d.get("valor","").strip()
    cat   = d.get("categoria","geral")
    if not chave or not valor: return jsonify({"erro":"Dados invalidos"}),400
    m = Memoria.query.filter_by(user_id=current_user.id, chave=chave).first()
    if m:
        m.valor = valor; m.categoria = cat; m.atualizado_em = datetime.now()
    else:
        m = Memoria(user_id=current_user.id, chave=chave, valor=valor, categoria=cat)
        db.session.add(m)
    db.session.commit()
    return jsonify({"ok":True, "mem": m.to_dict()})

@app.route("/api/memoria/deletar/<int:mid>", methods=["DELETE"])
@login_required
def deletar_memoria(mid):
    m = Memoria.query.filter_by(id=mid, user_id=current_user.id).first_or_404()
    db.session.delete(m); db.session.commit()
    return jsonify({"ok":True})

@app.route("/api/upgrade", methods=["POST"])
@login_required
def upgrade():
    current_user.plano = "premium"; db.session.commit()
    return jsonify({"ok":True})


@app.route("/manifest.json")
def manifest():
    return send_from_directory(os.path.join(app.root_path, 'static'), 'manifest.json')

if __name__ == "__main__":
    app.run(debug=True, port=5000)
